import json
import os
import requests
from docx import Document
import re

def get_valid_ip_port():
    """Reads IP and Port from settings.json."""
    current_dir = os.getcwd()
    SETTINGS_FILE = os.path.join(current_dir, "settings.json")
    
    with open(SETTINGS_FILE, "r") as file:
        settings = json.load(file)
    
    ip_settings = settings.get('ip_settings', {})
    ip = ip_settings.get('host')
    port = str(ip_settings.get('port'))
    
    return [ip, port]

def clean_text(text):
    """Remove special characters, brackets, quotes, and newlines."""
    cleaned = re.sub(r'[{}\[\]":\'"\n]', '', str(text))
    # Remove extra whitespaces
    cleaned = ' '.join(cleaned.split())
    return cleaned

def convert_json_to_docx(json_data, output_path):
    """Convert JSON data to DOCX format."""
    doc = Document()
    
    # Step 4: Write cleaned JSON content into the Word document
    if isinstance(json_data, dict):  # If JSON is a dictionary
        for key, value in json_data.items():
            cleaned_key = clean_text(key)
            cleaned_value = clean_text(value)
            doc.add_paragraph(f"{cleaned_key} {cleaned_value}")
    elif isinstance(json_data, list):  # If JSON is a list of dictionaries
        for item in json_data:
            if isinstance(item, dict):
                paragraph_text = []
                for key, value in item.items():
                    cleaned_key = clean_text(key)
                    cleaned_value = clean_text(value)
                    paragraph_text.append(f"{cleaned_key} {cleaned_value}")
                doc.add_paragraph(" ".join(paragraph_text))
                doc.add_paragraph("\n")  # Add space between entries
    
    doc.save(output_path)
    return True

def fetch_ddi_data(json_data=None):
    """Fetch drug-drug interaction data from the API."""
    try:
        # Get IP and port from settings
        ip, port = get_valid_ip_port()
        if not ip or not port:
            raise ValueError("IP address or port is missing. Please check your settings.")
            
        api_url = f"http://{ip}:{port}/uploadfile/"
        
        # Get current directory and file paths
        current_dir = os.getcwd()
        docx_file_path = os.path.join(current_dir, "RESOURCES", "converted_document_from_current_format_json.docx")
        
        # Use provided JSON data or read from file
        if json_data is None:
            json_file_path = os.path.join(current_dir, "RESOURCES", "current_format.json")
            with open(json_file_path, 'r') as f:
                json_data = json.load(f)
        
        # Convert JSON to DOCX
        convert_json_to_docx(json_data, docx_file_path)
        
        # Prepare files for API request
        files = {
            'file': ('document.docx', open(docx_file_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        # Make API request
        response = requests.post(api_url, files=files)
        files['file'][1].close()
        
        if response.status_code == 200:
            ddi_data = response.json()

            print("/n/n/n/n")
            print(ddi_data)
            print("/n/n/n/n")
            table_data = []
            if isinstance(ddi_data, dict):
                if 'interactions' in ddi_data:
                    for interaction in ddi_data['interactions']:
                        table_data.append({
                            'Drug 1': interaction.get('drug_A', ''),
                            'Drug 2': interaction.get('drug_B', ''),
                            'Interaction': interaction.get('interaction', '')
                        })
                elif 'drug_interactions' in ddi_data:
                    for interaction in ddi_data['drug_interactions']:
                        table_data.append({
                            'Drug 1': interaction.get('drug1', ''),
                            'Drug 2': interaction.get('drug2', ''),
                            'Interaction': interaction.get('description', '')
                        })
            return {'table': table_data}
        else:
            raise requests.exceptions.RequestException(f"Server returned status code {response.status_code}")
            
    except ConnectionRefusedError:
        raise ConnectionRefusedError("Connection refused. Please check if the DDI server is running and the IP/Port settings are correct.")
    except requests.exceptions.RequestException as e:
        if 'No host supplied' in str(e):
            raise ValueError("IP address is missing. Please check your IP settings in the Settings menu.")
        elif 'Failed to connect' in str(e):
            raise ConnectionError("Failed to connect to DDI server. Please check your IP and Port settings.")
        else:
            raise ConnectionError(f"Connection error: {str(e)}")
    except Exception as e:
        raise Exception(f"Error processing drug interactions: {str(e)}")

if __name__ == "__main__":
    # Dummy data for testing
    dummy_json_data = {
        "uuid": "ddb19471-364d-4e32-ac66-39e7055e5a24",
        "datetime": "23-03-2025 14:49:56",
        "date": "23-03-2025",
        "Name": "Test Patient",
        "Age_year": "78",
        "Age_month": "98",
        "Sex": "Male",
        "uhid": "UHID/123/ABC",
        "bed_number": "4",
        "Diagnosis": "Test Diagnosis",
        "Consultants": "Dr. One",
        "JR": "Dr. Three",
        "SR": "Dr. Four",
        "each_entry_layout": {
            "entry_1": {
                "title": "Respiratory support",
                "subtitles": {
                    "subtitle_1": {
                        "content": "O2 via NC",
                        "day": "D3",
                        "dose": "2L",
                        "volume": "N/A"
                    }
                }
            },
            "entry_2": {
                "title": "Antimicrobials",
                "subtitles": {
                    "subtitle_1": {
                        "content": "Amoxicillin",
                        "day": "D1",
                        "dose": "500mg",
                        "volume": "5ml"
                    },
                    "subtitle_2": {
                        "content": "Azithromycin",
                        "day": "D2",
                        "dose": "250mg",
                        "volume": "3ml"
                    }
                }
            },
            "entry_3": {
                "title": "Other Medications",
                "subtitles": {
                    "subtitle_1": {
                        "content": "Paracetamol",
                        "day": "D1",
                        "dose": "650mg",
                        "volume": "10ml"
                    },
                    "subtitle_2": {
                        "content": "Ibuprofen",
                        "day": "D2",
                        "dose": "400mg",
                        "volume": "8ml"
                    }
                }
            }
        }
    }

    # Test the conversion and API call with dummy data
    ddi_data = fetch_ddi_data(dummy_json_data)
    print(json.dumps(ddi_data, indent=2)) 
import sys
from PyQt5 import QtWidgets, QtGui, QtCore
from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QTextEdit,
    QTableWidget, QTableWidgetItem, QHeaderView, QMainWindow, QShortcut,
    QInputDialog, QCalendarWidget, QMessageBox, QLineEdit, QFrame,
    QSizePolicy, QApplication
)
from PyQt5.QtGui import QColor

import json
import requests
from docx import Document
import os

import os
import re

current_dir = os.getcwd()





def catch_exceptions(handler=None):
    """
    Decorator to catch exceptions in the decorated function.
    :param handler: Optional function to handle the exception.
    """

    def decorator(func):
        def wrapper(*args, **kwargs):
            try:
                print(f" [======]in {func.__name__}:")
                return func(*args, **kwargs)
            except Exception as e:
                if handler:
                    # Call the custom handler if provided
                    handler(e)
                else:
                    # Default behavior: Print the exception
                    print(args, kwargs)
                    print(f"An error occurred in {func.__name__}: {e}")

                    #####################################
                    ###########--ERROR BOX---############
                    error_box = QMessageBox()
                    error_box.setIcon(QMessageBox.Critical)  # Set error icon
                    error_box.setWindowTitle("Error")
                    error_box.setText(f"An unexpected error occurred! in {func.__name__}")
                    error_box.setInformativeText(f"{e}")
                    error_box.setStandardButtons(QMessageBox.Ok)

                    # Show the error box
                    error_box.exec_()
                    #####################################
                    ###########--ERROR BOX---############

        return wrapper

    return decorator

@catch_exceptions()
def get_valid_ip_port():
    """Reads, verifies, and extracts IP and Port from settings.txt.

    Returns:
        list: [ip, port] if valid.

    Raises:
        ValueError: If IP or Port is missing or has an invalid format.
        FileNotFoundError: If settings.txt does not exist.
    """
    SETTINGS_FILE = current_dir+"/RESOURCES/ipsettings.txt"

    # Ensure the settings file exists
    if not os.path.exists(SETTINGS_FILE):
        raise FileNotFoundError(f"Settings file '{SETTINGS_FILE}' not found.")

    ip = None
    port = None

    # Read and validate settings
    with open(SETTINGS_FILE, "r") as file:
        for line in file:
            if line.startswith("ip="):
                potential_ip = line.strip().split("=")[-1]
                ip_regex = r"^(?:\d{1,3}\.){3}\d{1,3}$"
                if re.match(ip_regex, potential_ip):
                    ip = potential_ip
                else:
                    raise ValueError(f"Invalid IP format: {potential_ip}")

            elif line.startswith("port="):
                potential_port = line.strip().split("=")[-1]
                if potential_port.isdigit() and 1 <= int(potential_port) <= 65535:
                    port = potential_port
                else:
                    raise ValueError(f"Invalid Port format: {potential_port}")

    # Ensure both IP and Port are found
    if ip is None or port is None:
        raise ValueError("Missing IP or Port in settings file.")

    return [ip, port]


class DDIWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Drug-Drug Interaction Scanner")
        self.resize(800, 600)

        # Main layout
        main_layout = QVBoxLayout(self)

        # Header section
        header_layout = QHBoxLayout()
        header_label = QLabel("DDI Scanner")
        header_label.setStyleSheet("font-size: 18px; font-weight: bold;")
        header_layout.addWidget(header_label)
        header_layout.addStretch()

        # Results section
        self.results_table = QTableWidget(0, 3)
        self.results_table.setHorizontalHeaderLabels(["Drug A", "Drug B", "Interaction"])
        self.results_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)

        # Add all sections to main layout
        main_layout.addLayout(header_layout)
        main_layout.addWidget(QLabel("Potential Interactions:"))
        main_layout.addWidget(self.results_table)

        # Style the window to match main app
        self.setStyleSheet("""
            QDialog {
                background-color: #D4D0C8;
            }
            QPushButton {
                background-color: #D4D0C8;
                border: 1px solid #808080;
                border-style: solid;
                padding: 4px 8px;
            }
            QPushButton:pressed {
                background-color: #C0C0C0;
                border: 1px solid #404040;
            }
            QTextEdit {
                background-color: white;
                border: 1px solid #808080;
            }
            QTableWidget {
                background-color: white;
                gridline-color: #808080;
                border: 1px solid #808080;
            }
        """)

        self.perform_scan()

    @catch_exceptions()
    def perform_scan(self):
        # Clear existing results
        self.results_table.setRowCount(0)

        try:
            # Call the API and get the response
            print("USING API TO FETCH DATA\n\n")
            response = self.fetch_api()

            # Check if we have a valid response
            if not response:
                QMessageBox.warning(self, "No Data", "Could not retrieve drug interaction data.")
                return

            # Extract interactions from the response
            interactions = response.get('interactions', [])

            # Populate table with interaction results
            for i, interaction in enumerate(interactions):
                drug_a = interaction.get('drug_A', '')
                drug_b = interaction.get('drug_B', '')
                interaction_text = interaction.get('interaction', '')

                # Use the interaction text as severity
                severity = str(interaction_text).lower()
                print(f"Severity for {drug_a} and {drug_b}: {severity}")  # Debug print

                self.results_table.insertRow(i)

                # Create table items
                item_a = QTableWidgetItem(drug_a)
                item_b = QTableWidgetItem(drug_b)
                item_interaction = QTableWidgetItem(interaction_text)

                # Color setting with broader matching
                if 'major' in severity:
                    color = QColor(255, 200, 200)  # Light red
                elif 'moderate' in severity:
                    color = QColor(255, 255, 200)  # Light yellow
                elif 'minor' in severity:
                    color = QColor(200, 255, 200)  # Light green
                else:
                    color = QColor(240, 240, 240)  # Light gray for unknown

                # Apply color to all items
                for item in [item_a, item_b, item_interaction]:
                    item.setBackground(color)
                    item.setForeground(QColor(0, 0, 0))  # Ensure black text

                # Add items to the table
                self.results_table.setItem(i, 0, item_a)
                self.results_table.setItem(i, 1, item_b)
                self.results_table.setItem(i, 2, item_interaction)

            # Resize rows to content
            self.results_table.resizeRowsToContents()

        except Exception as e:
            print(f"Error processing scan results: {e}")
            QMessageBox.warning(self, "Error", f"Error processing scan results: {e}")

    def get_executable_path(self):
        """Returns the path where the executable is located."""
        if getattr(sys, 'frozen', False):  # Running as a PyInstaller bundle
            return os.path.dirname(sys.executable)
        return os.path.dirname(os.path.abspath(__file__))

    @catch_exceptions()
    def fetch_api(self):
        # Handle PyInstaller temporary directory
        base_path = self.get_executable_path()

        # Define filenames
        json_filename = "/RESOURCES/current_format.json"
        docx_filename = "/RESOURCES/converted_document_from_current_format_json.docx"

        # Construct full paths
        # json_file_path = os.path.join(base_path, json_filename)
        json_file_path = current_dir+json_filename
        # docx_file_path = os.path.join(base_path, docx_filename)
        docx_file_path = current_dir+docx_filename

        # Step 1: Check if JSON file exists
        if not os.path.exists(json_file_path):
            print(current_dir)
            print(f"Error: JSON file '{json_file_path}' not found. Current dir: {base_path}")
            return None

        # Delete existing docx file if it exists
        if os.path.exists(docx_file_path):
            try:
                os.remove(docx_file_path)
                print(f"Existing document '{docx_filename}' has been deleted.")
            except Exception as e:
                print(f"Error deleting existing document: {e}")
                return None

        # Step 2: Read JSON file
        with open(json_file_path, 'r', encoding='utf-8') as json_file:
            data = json.load(json_file)

        # Step 3: Create a new Word document
        doc = Document()

        # Helper function to clean text
        def clean_text(text):
            # Remove special characters, brackets, quotes, and newlines
            import re
            cleaned = re.sub(r'[{}\[\]":\'"\n]', '', str(text))
            # Remove extra whitespaces
            cleaned = ' '.join(cleaned.split())
            return cleaned

        # Step 4: Write cleaned JSON content into the Word document
        if isinstance(data, dict):  # If JSON is a dictionary
            for key, value in data.items():
                cleaned_key = clean_text(key)
                cleaned_value = clean_text(value)
                doc.add_paragraph(f"{cleaned_key} {cleaned_value}")
        elif isinstance(data, list):  # If JSON is a list of dictionaries
            for item in data:
                if isinstance(item, dict):
                    paragraph_text = []
                    for key, value in item.items():
                        cleaned_key = clean_text(key)
                        cleaned_value = clean_text(value)
                        paragraph_text.append(f"{cleaned_key} {cleaned_value}")
                    doc.add_paragraph(" ".join(paragraph_text))
                    doc.add_paragraph("\n")  # Add space between entries

        # Step 5: Save the document
        doc.save(docx_file_path)
        print(f"Document '{docx_filename}' has been created.")

        ip_and_port = get_valid_ip_port()
        # Step 6: Send the document to API
        url = f'http://{ip_and_port[0]}:{ip_and_port[1]}/uploadfile/'

        with open(docx_file_path, 'rb') as f:
            files = {'file': (
                docx_filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}

            response = requests.post(url, files=files)

            # Step 7: Print response
            if response.status_code == 200:
                print("Response from the API:")
                print(response.json())  # Print JSON response
                return response.json()
            else:
                print(f"Error: {response.status_code}")
                print(response.text)
                return None


# For testing the dialog independently
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = DDIWindow()
    window.show()
    sys.exit(app.exec_())